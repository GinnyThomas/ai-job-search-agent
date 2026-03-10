from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(tailored: dict, candidate_name: str) -> bytes:
    """
    Convert the structured tailored content dict into a downloadable .docx.

    Returns bytes so app.py can pass it directly to st.download_button
    without writing a file to disk.

    Structure:
        Name (large, bold)
        Contact info (centred, small)
        ── Summary
        ── Key Skills
        ── Experience (variable bullets per role)
        ── Personal Projects
        ── Education
        ── Certifications & Development

    Cover note is intentionally excluded — it lives in the UI only,
    not on the CV itself.
    """
    doc = Document()

    # ── Remove default margins — use tighter CV margins ──────────────────
    section = doc.sections[0]
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    # ── Name ──────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(candidate_name)
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Contact info ──────────────────────────────────────────────────────
    contact_info = tailored.get("contact_info", "")
    if contact_info:
        contact_para = doc.add_paragraph(contact_info)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.runs[0].font.size = Pt(9)
        contact_para.paragraph_format.space_after = Pt(4)

    def add_section_heading(text):
        para = doc.add_paragraph()
        run = para.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        # Horizontal rule effect via bottom border
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
        return para

    def add_bullet(text):
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(text).font.size = Pt(10)
        para.paragraph_format.space_after = Pt(2)

    # ── Summary ───────────────────────────────────────────────────────────
    if tailored.get("summary"):
        add_section_heading("Summary")
        summary_para = doc.add_paragraph(tailored["summary"])
        summary_para.runs[0].font.size = Pt(10)

    # ── Skills ────────────────────────────────────────────────────────────
    skills = tailored.get("highlighted_skills", [])
    if skills:
        add_section_heading("Key Skills")
        skills_para = doc.add_paragraph(", ".join(skills))
        skills_para.runs[0].font.size = Pt(10)

    # ── Experience ────────────────────────────────────────────────────────
    experience = tailored.get("experience", [])
    if experience:
        add_section_heading("Experience")
        for role in experience:
            # Role title + dates on same line, right-aligned date
            role_para = doc.add_paragraph()
            role_para.paragraph_format.space_after = Pt(2)
            title_run = role_para.add_run(
                f"{role.get('role', '')} — {role.get('company', '')}"
            )
            title_run.bold = True
            title_run.font.size = Pt(10)
            dates = role.get("dates", "")
            if dates:
                role_para.add_run(f"  {dates}").font.size = Pt(9)

            for bullet in role.get("bullets", []):
                add_bullet(bullet)

    # ── Personal Projects ─────────────────────────────────────────────────
    projects = tailored.get("personal_projects", [])
    if projects:
        add_section_heading("Personal Projects")
        for project in projects:
            proj_para = doc.add_paragraph()
            proj_para.paragraph_format.space_after = Pt(2)
            proj_run = proj_para.add_run(project.get("name", ""))
            proj_run.bold = True
            proj_run.font.size = Pt(10)
            for bullet in project.get("bullets", []):
                add_bullet(bullet)

    # ── Education ─────────────────────────────────────────────────────────
    education = tailored.get("education", [])
    if education:
        add_section_heading("Education")
        for item in education:
            edu_para = doc.add_paragraph(item)
            edu_para.runs[0].font.size = Pt(10)
            edu_para.paragraph_format.space_after = Pt(2)

    # ── Certifications ────────────────────────────────────────────────────
    certifications = tailored.get("certifications", [])
    if certifications:
        add_section_heading("Certifications & Development")
        for item in certifications:
            cert_para = doc.add_paragraph(item)
            cert_para.runs[0].font.size = Pt(10)
            cert_para.paragraph_format.space_after = Pt(2)

    # ── Serialise to bytes ────────────────────────────────────────────────
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
